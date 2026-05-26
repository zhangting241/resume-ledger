"""
简历解析模块 - 支持 PDF 和 DOCX 格式
提取：姓名、年龄（从出生日期推算）、学历、工作经历（起止时间+公司+岗位）
兼容自由格式中文简历（大字标题姓名、三列式工作经历等）
"""
import re
import os
import json
from datetime import datetime
from collections import OrderedDict

# 调试目录
DEBUG_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "debug")
os.makedirs(DEBUG_DIR, exist_ok=True)


def parse_resume(filepath: str) -> dict:
    """解析简历文件，返回结构化数据"""
    ext = os.path.splitext(filepath)[1].lower()
    text = ""
    try:
        if ext == ".pdf":
            text = _extract_text_from_pdf(filepath)
        elif ext in (".docx", ".doc"):
            text = _extract_text_from_docx(filepath)
        else:
            text = _extract_text_from_txt(filepath)
    except Exception as e:
        print(f"[ERROR] 无法读取文件 {filepath}: {e}")
        return None

    if not text or len(text.strip()) < 20:
        return None

    # 清理文本
    text = _clean_text(text)

    result = {}
    result["name"] = _extract_name(text)
    result["age"] = _extract_age(text)
    result["education"] = _extract_education(text)
    result["work_experiences"] = _extract_work_experiences(text)
    result["phone"] = _extract_phone(text)
    result["email"] = _extract_email(text)
    result["gender"] = _extract_gender(text)
    result["city"] = _extract_city(text)
    result["filename"] = os.path.basename(filepath)

    return result


def _extract_text_from_pdf(filepath: str) -> str:
    text = ""
    
    # 方案1: pdfplumber.extract_text()
    try:
        import pdfplumber
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
                # 如果 extract_text 返回空但页面有 chars，手动拼接
                if not page_text and page.chars:
                    chars_sorted = sorted(page.chars, key=lambda c: (-c['top'], c['x0']))
                    line_chars = []
                    current_top = None
                    for c in chars_sorted:
                        c_top = round(float(c['top']), 1)
                        if current_top is not None and abs(c_top - current_top) > 1.5:
                            line_chars.append('\n')
                        line_chars.append(c['text'])
                        current_top = c_top
                    if line_chars:
                        text += ''.join(line_chars) + "\n"
    except Exception as e:
        print(f"[WARN] pdfplumber 提取失败: {e}")

    # 方案2: 如果 pdfplumber 提取为空，用 pdfminer.six 作为备用
    if not text.strip():
        try:
            from pdfminer.high_level import extract_text as pdfminer_extract
            print("[INFO] 使用 pdfminer 备用方案提取...")
            text = pdfminer_extract(filepath)
        except Exception as e:
            print(f"[WARN] pdfminer 提取失败: {e}")
    
    # 方案3: 最后尝试 PyPDF2
    if not text.strip():
        try:
            from PyPDF2 import PdfReader
            print("[INFO] 使用 PyPDF2 备用方案提取...")
            reader = PdfReader(filepath)
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    text += t + "\n"
        except Exception as e:
            print(f"[WARN] PyPDF2 提取失败: {e}")

    # 保存调试文本
    _src_filename = os.path.basename(filepath)
    debug_path = os.path.join(DEBUG_DIR, f"{_src_filename}.txt")
    try:
        with open(debug_path, "w", encoding="utf-8") as f:
            f.write(f"=== 原始文本 ({len(text)} 字符) ===\n")
            f.write(text)
            f.write("\n\n=== 文本统计 ===\n")
            f.write(f"行数: {len(text.splitlines())}\n")
            has_chinese = [l for l in text.splitlines() if any('\u4e00' <= c <= '\u9fff' for c in l)]
            f.write(f"含中文行数: {len(has_chinese)}\n")
            if has_chinese:
                f.write("前10个中文行:\n")
                for l in has_chinese[:10]:
                    f.write(f"  {l}\n")
    except Exception as e:
        print(f"[DEBUG] 保存调试文本失败: {e}")

    print(f"[DEBUG] PDF 提取文本长度: {len(text)} 字符 -> {debug_path}")
    return text


def _extract_text_from_docx(filepath: str) -> str:
    from docx import Document
    doc = Document(filepath)
    lines = []
    for para in doc.paragraphs:
        lines.append(para.text)
    # 也读取表格
    for table in doc.tables:
        for row in table.rows:
            row_text = "  ".join(cell.text for cell in row.cells)
            lines.append(row_text)
    return "\n".join(lines)


def _extract_text_from_txt(filepath: str) -> str:
    with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def _clean_text(text: str) -> str:
    """清理多余空白，保留换行结构和多空格列分隔符"""
    # 统一换行符
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    # 压缩连续空行为单个空行
    text = re.sub(r'\n{3,}', '\n\n', text)
    # 行内Tab转空格（保留列分隔效果）
    text = text.replace('\t', '    ')
    # 压缩超长空格序列（>4个）但保留2-4个空格作为列分隔线索
    text = re.sub(r'[ ]{5,}', '    ', text)
    return text.strip()


def _extract_name(text: str) -> str:
    """
    提取姓名：
    1. 先尝试 "姓名：XXX" 标准格式
    2. 在前20行中找首个 2-4 个连续中文字符（排除常见非姓名词汇）
    3. 放宽到前800字符中的汉字组合
    """
    # 方式1: 标准前缀
    patterns = [
        r'姓名[：:\s]*([^\s,，。；;]{2,4})',
        r'[Nn]ame[：:\s]*([^\s,，。；;]{2,20})',
        r'姓\s*名[：:\s]*([^\s,，。；;]{2,4})',
    ]
    for p in patterns:
        m = re.search(p, text[:800])
        if m:
            name = m.group(1).strip()
            if re.match(r'^[\u4e00-\u9fff·]{2,4}$', name):
                return name

    # 方式2: 扫描前20行，找独立的2-4个汉字
    lines = text.strip().split('\n')
    skip_keywords = ['简历', '个人', '应聘', '求职', '联系', '电话', '邮箱', '地址',
                     '性别', '民族', '政治', '出生', '现住', '教育', '工作',
                     '自我评价', '项目经历', '技能证书', '基本信息', '免冠']
    for line in lines[:20]:
        line = line.strip()
        if not line or len(line) > 20:
            continue
        if any(kw in line for kw in skip_keywords):
            continue
        # 独立行 2-4 个汉字
        m = re.search(r'^([\u4e00-\u9fff·]{2,4})$', line)
        if m:
            return m.group(1)

    # 方式3: 更宽松——在前800字符中找最前面的2-4个连续汉字（不在句子中）
    prefix = text[:800]
    # 去掉标点后找
    clean_prefix = re.sub(r'[|｜/\\,:：;；.。，、()（）【】\[\]\s\d]', ' ', prefix)
    tokens = clean_prefix.split()
    for token in tokens:
        token = token.strip()
        if re.match(r'^[\u4e00-\u9fff·]{2,4}$', token):
            if not any(kw in token for kw in skip_keywords):
                return token

    return "未知"


def _extract_age(text: str) -> str:
    """从出生日期推算年龄，或直接匹配年龄"""
    # 先尝试直接匹配年龄
    for p in [r'年龄[：:\s]*(\d{1,2})\s*岁?', r'(\d{1,2})\s*岁']:
        m = re.search(p, text[:500])
        if m:
            age = int(m.group(1))
            if 18 <= age <= 70:
                return str(age)

    # 出生年月/日期模式
    birth_patterns = [
        r'(?:出生|生日)[日期年月]?[：:\s]*(\d{4})[年/\-.](\d{1,2})[月/\-.]?(\d{1,2})?',
        r'(\d{4})[年/\-.](\d{1,2})[月.\-](\d{1,2})[日出]?生',
        r'出生[于在]?[：:\s]*(\d{4})[年/\-.](\d{1,2})[月.\-]?(\d{1,2})?',
        r'(?:出生年月|出生日期)[：:\s]*(\d{4})[年/\-.](\d{1,2})',
    ]

    for p in birth_patterns:
        m = re.search(p, text[:500])
        if m:
            year = int(m.group(1))
            month = int(m.group(2)) if m.lastindex >= 2 else 6
            day = int(m.group(3)) if m.lastindex >= 3 and m.group(3) else 15
            try:
                birth = datetime(year, month, day)
                today = datetime.now()
                age = today.year - birth.year - ((today.month, today.day) < (birth.month, birth.day))
                if 18 <= age <= 70:
                    return str(age)
            except:
                pass

    return "未知"


def _extract_education(text: str) -> str:
    """提取最高学历，优先从教育经历段落识别"""
    edu_map = OrderedDict([
        ("博士研究生", ["博士研究生", "博士"]),
        ("硕士研究生", ["硕士研究生", "硕士"]),
        ("本科", ["本科", "学士", "大学本科"]),
        ("大专", ["大专", "专科", "高职"]),
        ("高中/中专", ["高中", "中专", "中技", "职高"]),
    ])

    # 策略1: "学历："关键词
    m = re.search(r'学历[：:\s]*([^\s,，。；;]{2,20})', text[:600])
    if m:
        edu_text = m.group(1).strip()
        for level, keywords in edu_map.items():
            for kw in keywords:
                if kw in edu_text:
                    return level

    # 策略2: 从教育经历段落提取
    edu_section = _find_section(text, [r'教\s*育\s*经\s*历', r'教\s*育\s*背\s*景',
                                       r'学\s*习\s*经\s*历', r'Education'])
    if edu_section:
        # 教育段落中，按优先级匹配学历关键词
        for level, keywords in edu_map.items():
            for kw in keywords:
                if kw in edu_section:
                    return level

    # 策略3: 全文搜索（限制前1500字符避免尾部培训信息干扰）
    text_upper = text[:1500]
    for level, keywords in edu_map.items():
        for kw in keywords:
            if kw in text_upper:
                return level

    return "未知"


def _extract_phone(text: str) -> str:
    """提取手机号"""
    m = re.search(r'1[3-9]\d{9}', text[:500])
    return m.group(0) if m else ""


def _extract_email(text: str) -> str:
    """提取邮箱"""
    m = re.search(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text[:500])
    return m.group(0) if m else ""


def _extract_gender(text: str) -> str:
    """提取性别"""
    m = re.search(r'性别[：:]\s*([男女])\s*', text[:500])
    if m:
        return m.group(1)
    m = re.search(r'[，,\s]([男女])\s*(?:[，,性族])', text[:500])
    if m:
        return m.group(1)
    return ""


def _extract_city(text: str) -> str:
    """提取所在城市"""
    patterns = [
        r'(?:现住城市|所在城市|现居|城市|居住地|所在地)[：:]\s*([^\s\n]{2,10})',
    ]
    for p in patterns:
        m = re.search(p, text[:800])
        if m:
            city = m.group(1).strip().rstrip('，,。.')
            # 过滤掉非城市的内容
            if len(city) >= 2 and not re.match(r'^[\d\.\-]+$', city):
                return city
    return ""


def _extract_work_experiences(text: str) -> list:
    """
    提取工作经历：精准定位"工作经历"标题 → 逐行解析下方内容
    支持格式：
      - 三列式自由格式：公司    岗位    2023.12-至今
      - 智联/Boss：公司 岗位（前一行）+ 时间（下一行）
      - 单行紧凑：时间 | 公司 | 岗位（任意顺序）
    """
    experiences = []

    # ===== 步骤1: 精准定位"工作经历"段落 =====
    work_section = _find_work_section(text)
    if not work_section:
        # 退而求其次：用全文
        work_section = text

    lines = work_section.split('\n')
    # 跳过标题行（第一行通常是"工作经历"本身）
    content_lines = []
    found_title = False
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if not found_title:
            # 检测标题行
            if re.search(r'工\s*作\s*经\s*历|工\s*作\s*经\s*验|职\s*业\s*经\s*历|从\s*业\s*经\s*历|Work\s*Experience|Professional\s*Experience|Employment', stripped, re.IGNORECASE):
                found_title = True
                continue
        if found_title:
            # 遇到下一个大标题就停止
            if _is_section_title(stripped):
                break
            content_lines.append(stripped)

    # ===== 步骤2: 逐行匹配时间范围 =====
    # 时间模式：2023.12-至今、2021.03-2023.11、2020年6月-2022年8月、2023/12-至今
    time_range_pattern = re.compile(
        r'(\d{4}\s*[年./\-]\s*\d{1,2})\s*[-~—至到]\s*(\d{4}\s*[年./\-]\s*\d{1,2}|至今|现在|目前|今)'
    )

    i = 0
    while i < len(content_lines):
        line = content_lines[i]
        if len(line) < 6:
            i += 1
            continue

        # 跳过纯描述行（以 • - · 开头）
        if re.match(r'^[•·\-–—●◆▪▸►]\s', line):
            i += 1
            continue

        m = time_range_pattern.search(line)
        if not m:
            # 该行没有时间 → 可能是只有公司+岗位的行（下一行有时间）
            i += 1
            continue

        start_time = m.group(1).replace(' ', '')
        end_time = m.group(2).replace(' ', '')

        # ===== 步骤3: 从这一行提取公司+岗位 =====
        company = ""
        position = ""

        # 3a) 时间前面的文本 = 公司 + 岗位
        before_time = line[:m.start()].strip()
        if before_time:
            company, position = _split_company_position(before_time)

        # 3b) 如果时间前面没有公司，检查上一行
        if not company and i > 0:
            prev = content_lines[i - 1]
            if prev and not _is_section_title(prev) and not time_range_pattern.search(prev):
                company, position = _split_company_position(prev)

        # 3c) 如果还没有，检查再上一行（智联格式：公司+岗位 | 空行 | 时间）
        if not company and i > 1:
            prev2 = content_lines[i - 2]
            if prev2 and not _is_section_title(prev2) and not time_range_pattern.search(prev2):
                company, position = _split_company_position(prev2)

        # 3d) 时间后面的文本作为补充
        if not company:
            after_time = line[m.end():].strip()
            after_time = _clean_work_rest(after_time)
            if after_time and len(after_time) >= 3:
                company, position = _split_company_position(after_time)

        if company and len(company) > 1:
            experiences.append({
                "period": f"{start_time}-{end_time}",
                "company": company,
                "position": position if position else "未识别"
            })

        i += 1

    return experiences


def _clean_work_rest(text: str) -> str:
    """清理工作经历剩余文本"""
    text = re.sub(r'^\s*(?:[\d一二三四五六七八九十]+[、\.。,，\)\]）])\s*', '', text)
    text = re.sub(r'^\s*(?:[\(（]\d+[\)）])\s*', '', text)
    text = re.sub(r'^\s*(?:[①②③④⑤⑥⑦⑧⑨⑩])\s*', '', text)
    text = re.sub(r'^\s*[\(（][^\)）]*[\)）]\s*', '', text)
    return text.strip()


def _is_section_title(text: str) -> bool:
    """判断是否为段落标题（不应被当作公司名）"""
    keywords = ['工作描述', '工作详情', '项目经历', '教育经历', '自我评价',
                '所获证书', '技能证书', '培训经历', '项目描述', '工作经历',
                '求职意向', '个人优势', '基本信息', '教育背景']
    return any(kw in text for kw in keywords)


def _find_section(text: str, patterns: list) -> str:
    """通用段落查找"""
    for p in patterns:
        m = re.search(p, text, re.IGNORECASE)
        if m:
            start = m.start()
            section = text[start:start + 3000]
            # 截断到下一个大标题
            stop_patterns = [
                r'\n\s*(项\s*目\s*经\s*历|教\s*育\s*背\s*景|教\s*育\s*经\s*历|自\s*我\s*评\s*价|技\s*能\s*证\s*书|培\s*训\s*经\s*历|工\s*作\s*经\s*历)',
                r'\n\s*(Project|Education|Skills|Certification|Self-evaluation|Work)',
            ]
            for sp in stop_patterns:
                sm = re.search(sp, section[20:], re.IGNORECASE)
                if sm:
                    section = section[:20 + sm.start()]
                    break
            return section
    return ""


def _find_work_section(text: str) -> str:
    """找到工作经历相关段落"""
    return _find_section(text, [
        r'工\s*作\s*经\s*历|工\s*作\s*经\s*验|职\s*业\s*经\s*历|从\s*业\s*经\s*历',
        r'Work\s*Experience|Professional\s*Experience|Employment',
    ])


def _split_company_position(text: str) -> tuple:
    """
    将工作经历描述拆分为 公司名 + 岗位
    核心策略（按优先级）：
      0. 多空格/Tab 列分隔（三列式简历：公司    岗位）
      1. 常见句式（就职于XX担任XX / XX-XX / XX|XX）
      2. 岗位关键词反向匹配（按长度降序，优先匹配"理财经理"而非"经理"）
      3. 单空格分割（右起第一个空格）
    """
    text = text.strip()
    if not text:
        return "", ""

    # ═══ 策略0: 多空格/Tab 列分隔（最高优先级，_clean_text 现已保留2-4空格） ═══
    if '  ' in text:
        # 用2个以上连续空格/制表符分割
        parts = re.split(r' {2,}|\t', text)
        parts = [p.strip() for p in parts if p.strip()]
        if len(parts) >= 3:
            # 三列及以上：公司  岗位  其他 → 取前两个
            return parts[0][:50], parts[1][:20]
        if len(parts) == 2:
            if _looks_like_position(parts[1]):
                return parts[0][:50], parts[1][:20]
            if _looks_like_position(parts[0]):
                return parts[1][:50], parts[0][:20]
            if len(parts[0]) > len(parts[1]):
                return parts[0][:50], parts[1][:20]
            return parts[1][:50], parts[0][:20]

    # ═══ 策略1: 常见句式 ═══
    # A: 就职于[公司]担任[岗位]
    m = re.search(r'(?:就职于|曾在|于|在|入职|加入)\s*([^\s,，。；;]{2,30})(?:担任|任|做|为)\s*([^\s,，。；;]{2,15})', text)
    if m:
        return m.group(1).strip(), m.group(2).strip()

    # B: [公司]担任[岗位]一职
    m = re.search(r'([^\s,，。；;]{3,30})(?:担任|任)\s*([^\s,，。；;]{2,15})(?:一职|岗位|工作|职务|职位)', text)
    if m:
        return m.group(1).strip(), m.group(2).strip()

    # C: [公司]-[岗位]
    m = re.search(r'(.{3,30})\s*-\s*(.{2,20})$', text)
    if m:
        return m.group(1).strip(), m.group(2).strip()

    # D: [公司]|[岗位]
    m = re.search(r'(.{3,30})\s*[｜|]\s*(.{2,20})$', text)
    if m:
        return m.group(1).strip(), m.group(2).strip()

    # E: [岗位]于[公司]
    m = re.search(r'([^\s,，。；;]{2,15})(?:于|在)\s*([^\s,，。；;]{3,30})', text)
    if m:
        return m.group(2).strip(), m.group(1).strip()

    # ═══ 策略2: 岗位关键词反向匹配（按长度降序） ═══
    position_keywords = sorted([
        '理财经理', '大堂经理', '客户经理', '柜员', '信贷经理',
        '工程师', '经理', '主管', '总监', '专员', '助理', '顾问',
        '架构师', '程序员', '开发者', '运营', '产品', '设计师',
        '会计', '出纳', '行政', '人事', 'HR', '销售', '市场',
        '实习', '管培生', 'VP', 'CEO', 'CTO', 'COO', 'CFO',
        '总裁', '副总裁', '主任', '科长', '处长', '局长',
        '教授', '讲师', '研究员', '分析师', '策划', '编辑',
        '店长', '副店长', '店员', '客服', '品质', '安全',
        '管理', '物业', '招商', '保安', '保洁', '绿化工',
        '项目经理', '客服经理', '工程主管', '工程经理',
        '副经理', '区域经理', '城市经理', '业务经理',
        '安装工程师', '土建工程师', '电气工程师',
        '技术支持', '技术经理', '技术主管',
        '维修', '技工', '电工', '水暖工', '木工',
        '消防', '监控', '秩序', '巡逻',
        '管家', '楼栋', '客服前台',
        '中控', '中控员', '安全员', '施工员',
        'Manager', 'Engineer', 'Director', 'Lead',
        '开发', '测试', '前端', '后端', '全栈', '运维',
    ], key=len, reverse=True)

    for kw in position_keywords:
        idx = text.rfind(kw)
        if idx >= 2:
            company = text[:idx].strip().rstrip('，,。.担任任就职于在加入和及与、-|｜ ')
            position = text[idx:].strip().lstrip('，,。.')
            if len(company) >= 2 and len(position) >= 2:
                return company[:50], position[:20]

    # ═══ 策略3: 单空格分割（右起第一个空格） ═══
    if ' ' in text:
        idx = text.rfind(' ')
        if 2 < idx < len(text) - 1:
            left = text[:idx].strip()
            right = text[idx+1:].strip()
            if len(left) >= 2 and len(right) >= 2:
                if _looks_like_position(right) or len(right) <= 20:
                    return left[:50], right[:20]

    # ═══ 策略4: 短文本 — 末尾中文当岗位 ═══
    if 4 <= len(text) <= 30:
        m = re.match(r'([\u4e00-\u9fff（）()a-zA-Z·&]{3,20})([\u4e00-\u9fff]{2,15})$', text)
        if m:
            return m.group(1), m.group(2)

    # ═══ 最终兜底：含"担任" ═══
    if '担任' in text:
        idx = text.index('担任')
        if idx > 2:
            return text[:idx].strip()[:50], text[idx+2:].strip()[:20]

    return text[:50], "未识别"


def _looks_like_position(text: str) -> bool:
    """判断文本是否看起来像岗位名称"""
    position_keywords = [
        '工程师', '经理', '主管', '总监', '专员', '助理', '顾问',
        '架构师', '程序员', '开发者', '运营', '产品', '设计师',
        '会计', '出纳', '行政', '人事', 'HR', '销售', '市场',
        '实习', '管培生', '总裁', '副总裁', '主任', '科长',
        '教授', '讲师', '研究员', '分析师', '策划', '编辑',
        '店长', '副店长', '店员', '客服', '品质', '安全',
        '管理', '物业', '招商', '保安', '保洁', '绿化工',
        '项目经理', '客服经理', '工程主管', '工程经理',
        '副经理', '区域经理', '城市经理', '业务经理',
        '安装工程师', '土建工程师', '电气工程师',
        '技术支持', '技术经理', '技术主管',
        '维修', '技工', '电工', '水暖工', '木工',
        '消防', '监控', '秩序', '巡逻',
        '管家', '楼栋', '客服前台',
        '中控', '中控员', '安全员', '施工员',
        '理财经理', '大堂经理', '客户经理', '柜员',
        'Manager', 'Engineer', 'Director', 'Lead',
        '开发', '测试', '前端', '后端', '全栈', '运维',
    ]
    return any(kw in text for kw in position_keywords)


if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1:
        result = parse_resume(sys.argv[1])
        import json
        print(json.dumps(result, ensure_ascii=False, indent=2))
